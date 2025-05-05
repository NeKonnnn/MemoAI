import os
import tempfile
import docx
import PyPDF2
import openpyxl
import pdfplumber
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.docstore.document import Document

class DocumentProcessor:
    def __init__(self):
        # Инициализация векторного хранилища с пустым набором
        self.documents = []
        self.doc_names = []
        self.embeddings = None
        self.vectorstore = None
        self.init_embeddings()
        
    def init_embeddings(self):
        """Инициализация модели для эмбеддингов"""
        try:
            # Загружаем модель для русского языка
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
            )
            print("Модель эмбеддингов успешно загружена")
        except Exception as e:
            print(f"Ошибка при загрузке модели эмбеддингов: {str(e)}")
            self.embeddings = None
    
    def process_document(self, file_path):
        """Обработка документа в зависимости от его типа"""
        file_extension = os.path.splitext(file_path)[1].lower()
        document_text = ""
        
        try:
            if file_extension == '.docx':
                document_text = self.extract_text_from_docx(file_path)
            elif file_extension == '.pdf':
                document_text = self.extract_text_from_pdf(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                document_text = self.extract_text_from_excel(file_path)
            elif file_extension == '.txt':
                document_text = self.extract_text_from_txt(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.webp']:
                document_text = self.extract_text_from_image(file_path)
            else:
                return False, f"Неподдерживаемый формат файла: {file_extension}"
            
            # Добавляем документ в коллекцию
            self.add_document_to_collection(document_text, os.path.basename(file_path))
            return True, f"Документ {os.path.basename(file_path)} успешно обработан"
            
        except Exception as e:
            return False, f"Ошибка при обработке документа: {str(e)}"
    
    def extract_text_from_docx(self, file_path):
        """Извлечение текста из DOCX файла"""
        doc = docx.Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return "\n".join(full_text)
    
    def extract_text_from_pdf(self, file_path):
        """Извлечение текста из PDF файла"""
        text = ""
        
        # Используем PDFPlumber для более точного извлечения текста
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            print(f"Ошибка при извлечении текста с помощью pdfplumber: {str(e)}")
            
            # Резервный метод с PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        text += page.extract_text() or ""
            except Exception as e2:
                print(f"Ошибка при извлечении текста с помощью PyPDF2: {str(e2)}")
                raise
        
        return text
    
    def extract_text_from_excel(self, file_path):
        """Извлечение текста из Excel файла"""
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text_content = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_content.append(f"Лист: {sheet_name}")
            
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    text_content.append("\t".join(row_values))
        
        return "\n".join(text_content)
    
    def extract_text_from_txt(self, file_path):
        """Извлечение текста из TXT файла"""
        try:
            # Пробуем открыть файл как UTF-8
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Если не удалось открыть как UTF-8, пробуем другие кодировки
            encodings = ['cp1251', 'latin-1', 'koi8-r']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # Если все кодировки не подошли, открываем в бинарном режиме
            with open(file_path, 'rb') as file:
                content = file.read()
                return str(content)

    def extract_text_from_image(self, file_path):
        """Извлечение текста из изображения с помощью OCR"""
        try:
            # Проверяем наличие библиотеки pytesseract
            import pytesseract
            from PIL import Image
            
            # Открываем изображение с помощью Pillow
            img = Image.open(file_path)
            
            # Извлекаем текст с изображения
            text = pytesseract.image_to_string(img, lang='rus+eng')
            
            # Если текст не извлечен, добавляем описание изображения
            if not text.strip():
                return f"[Изображение: {os.path.basename(file_path)}. OCR не смог извлечь текст.]"
            
            return text
        except ImportError:
            # Если pytesseract не установлен, возвращаем информацию о файле
            return f"[Изображение: {os.path.basename(file_path)}. Для распознавания текста требуется установка pytesseract.]"
        except Exception as e:
            return f"[Изображение: {os.path.basename(file_path)}. Ошибка при обработке: {str(e)}]"
    
    def add_document_to_collection(self, text, doc_name):
        """Добавление документа в коллекцию и обновление векторного хранилища"""
        # Разбиваем текст на части
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        chunks = text_splitter.split_text(text)
        
        # Создаем документы для langchain
        langchain_docs = []
        for i, chunk in enumerate(chunks):
            langchain_docs.append(
                Document(
                    page_content=chunk,
                    metadata={"source": doc_name, "chunk": i}
                )
            )
        
        # Добавляем в общий список документов
        self.documents.extend(langchain_docs)
        if doc_name not in self.doc_names:
            self.doc_names.append(doc_name)
        
        # Обновляем векторное хранилище
        self.update_vectorstore()
    
    def update_vectorstore(self):
        """Обновление или создание векторного хранилища"""
        if not self.documents:
            print("Нет документов для индексации")
            return
        
        if not self.embeddings:
            print("Модель эмбеддингов не инициализирована")
            self.init_embeddings()
            if not self.embeddings:
                return
        
        try:
            # Создаем новое векторное хранилище
            self.vectorstore = FAISS.from_documents(self.documents, self.embeddings)
            print(f"Векторное хранилище обновлено, добавлено {len(self.documents)} чанков")
        except Exception as e:
            print(f"Ошибка при обновлении векторного хранилища: {str(e)}")
    
    def query_documents(self, query, k=5):
        """Поиск релевантных документов по запросу"""
        if not self.vectorstore:
            return "Векторное хранилище не инициализировано или пусто"
        
        try:
            docs = self.vectorstore.similarity_search(query, k=k)
            results = []
            
            for doc in docs:
                results.append({
                    "content": doc.page_content,
                    "source": doc.metadata.get("source", "Неизвестный источник"),
                    "chunk": doc.metadata.get("chunk", 0)
                })
            
            return results
        except Exception as e:
            return f"Ошибка при поиске по документам: {str(e)}"
    
    def get_document_list(self):
        """Получение списка загруженных документов"""
        return self.doc_names
    
    def clear_documents(self):
        """Очистка коллекции документов"""
        self.documents = []
        self.doc_names = []
        self.vectorstore = None
        return "Коллекция документов очищена"
    
    def process_query(self, query, agent_function):
        """Обработка запроса с контекстом документов для LLM"""
        if not self.vectorstore:
            return "Нет загруженных документов. Пожалуйста, загрузите документы перед выполнением запроса."
        
        try:
            # Получаем релевантные документы
            docs = self.query_documents(query)
            
            if isinstance(docs, str):  # Если возникла ошибка
                return docs
            
            # Формируем контекст из найденных документов
            context = "Контекст из документов:\n\n"
            for i, doc in enumerate(docs):
                context += f"Фрагмент {i+1} (из документа '{doc['source']}'):\n{doc['content']}\n\n"
            
            # Подготавливаем запрос для LLM с инструкциями и контекстом
            prompt = f"""На основе предоставленного контекста ответь на вопрос пользователя. 
Если информации в контексте недостаточно, укажи это.
Отвечай только на основе информации из контекста. Не придумывай информацию.

{context}

Вопрос пользователя: {query}

Ответ:"""
            
            # Отправляем запрос к LLM
            response = agent_function(prompt)
            return response
            
        except Exception as e:
            return f"Ошибка при обработке запроса: {str(e)}" 